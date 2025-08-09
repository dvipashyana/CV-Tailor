import re
import io
import os
import json
from datetime import datetime
from pathlib import Path
import numpy as np
import pandas as pd
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document

# ----------------------------- #
# Simple, light-weight helpers  #
# ----------------------------- #

APP_DIR = Path.cwd()
SAVED_DIR = APP_DIR / "saved_cvs"
META_FILE = SAVED_DIR / "index.json"
SAVED_DIR.mkdir(exist_ok=True)
if not META_FILE.exists():
    META_FILE.write_text(json.dumps([], indent=2))

STOPWORDS = set("""
a about above after again against all am an and any are aren't as at be because been
before being below between both but by can't cannot could couldn't did didn't do
does doesn't doing don't down during each few for from further had hadn't has hasn't
have haven't having he he'd he'll he's her here here's hers herself him himself his
how how's i i'd i'll i'm i've if in into is isn't it it's its itself let's me more
most mustn't my myself no nor not of off on once only or other ought our ours
ourselves out over own same shan't she she'd she'll she's should shouldn't so some such
than that that's the their theirs them themselves then there there's these they
they'd they'll they're they've this those through to too under until up very was wasn't we
we'd we'll we're we've were weren't what what's when when's where where's which while who
who's whom why why's with won't would wouldn't you you'd you'll you're you've your yours
yourself yourselves
""".split())

SENIORITY_TERMS = {"manager","lead","senior","principal","head","strategy","governance"}
IMPACT_TERMS = {"increase","reduce","improve","optimize","grow","accelerate","lower","raise",
                "decrease","streamline","automate","launch","deliver","scale","save"}
RESULT_TERMS = {"by","resulting","leading","thereby","which","so that"}

VERB_BANK = [
    "Drove","Led","Owned","Built","Launched","Optimized","Streamlined","Automated","Delivered",
    "Designed","Implemented","Developed","Standardized","Improved","Accelerated","Reduced","Increased"
]

def read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    return "\n".join(text).strip()

def write_docx(cv_title: str, tvps: str, sections: dict) -> bytes:
    doc = Document()
    doc.add_heading(cv_title, level=0)
    if tvps:
        doc.add_paragraph(tvps)
        doc.add_paragraph("")  # spacer
    for section, bullets in sections.items():
        if not bullets: 
            continue
        doc.add_heading(section, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def clean_tokens(text: str):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s\-\+\%\.]', ' ', text)
    toks = [t for t in text.split() if t not in STOPWORDS and len(t) > 2]
    return toks

def top_keywords(text: str, n=20):
    if not text or text.strip()=="":
        return []
    corpus = [text]
    from sklearn.feature_extraction.text import TfidfVectorizer
    vec = TfidfVectorizer(stop_words='english', ngram_range=(1,2), min_df=1)
    X = vec.fit_transform(corpus)
    scores = X.toarray()[0]
    items = sorted(zip(vec.get_feature_names_out(), scores), key=lambda x: x[1], reverse=True)
    # Filter noise
    filtered = []
    for term, score in items:
        if any(c.isalpha() for c in term) and not term.isdigit() and len(term) > 2:
            filtered.append((term, score))
    return [t for t,_ in filtered[:n]]

def extract_metrics(text: str):
    return re.findall(r"(\d+(?:\.\d+)?\s?%?)", text)

def xyz_rewrite(bullet: str, role_keywords=None, jd_outcomes=None):
    """
    Heuristic XYZ rewrite:
    - Try to detect Action (verb lead)
    - Identify mechanism (how)
    - Identify result (numbers/impact)
    """
    role_keywords = role_keywords or []
    jd_outcomes = jd_outcomes or []
    original = bullet.strip().rstrip(".")
    if not original:
        return ""
    # Capture any numbers as Z
    nums = extract_metrics(original)
    z_hint = nums[0] if nums else None

    # Choose a strong verb
    first_word = original.split()[0]
    verb = first_word if first_word[0].isupper() else None
    if verb is None or len(verb) < 3:
        verb = np.random.choice(VERB_BANK)

    # Try to find a 'by <how>' clause
    by_clause = ""
    m = re.search(r"\bby\b\s(.+)", original, flags=re.IGNORECASE)
    if m:
        by_clause = m.group(1)
    else:
        # fallback: take middle chunk as 'by ...'
        parts = re.split(r",|–|-", original)
        if len(parts) > 1:
            by_clause = parts[1].strip()
        else:
            by_clause = original

    # Try to pull an outcome term
    outcome_pool = list(IMPACT_TERMS.union(set(jd_outcomes)))
    outcome_word = np.random.choice(outcome_pool) if outcome_pool else "improved"

    # Optional role keyword to anchor relevance
    role_anchor = ""
    for kw in role_keywords:
        if len(kw) > 3 and " " not in kw:
            role_anchor = kw
            break

    # Build XYZ
    if z_hint:
        z_text = f"by {z_hint}"
    else:
        z_text = "— ⚠ add a measurable result (%, £, time saved, error rate, etc.)"

    anchor = f" for {role_anchor}" if role_anchor else ""
    result_phrase = f" resulting in {z_text}" if z_hint else f" {z_text}"

    return f"{verb} {outcome_word}{anchor} by {by_clause}{result_phrase}."

def hiring_manager_notes(jd_text: str):
    """
    Produce crisp guidance bullets from the JD, focusing on POV the manager cares about.
    """
    kws = top_keywords(jd_text, n=25)
    # Group buckets
    buckets = {
        "Must-haves": [],
        "Great-to-haves": [],
        "Impact to Prove": [],
    }
    for k in kws:
        if any(s in k for s in ["governance","taxonomy","quality","privacy","compliance","agile","api","sql","python","stakeholder","kpi","media","campaign","measurement"]):
            buckets["Must-haves"].append(k)
        elif any(s in k for s in ["cloud","ml","ai","dbt","bigquery","snowflake","ga4","clean room"]):
            buckets["Great-to-haves"].append(k)
        else:
            buckets["Impact to Prove"].append(k)
    return buckets

def tvps_generator(role_title: str, jd_text: str):
    # short, punchy 2–3 lines
    foc = top_keywords(jd_text, n=8)
    foc = [f for f in foc if f not in STOPWORDS][:6]
    foc_txt = " • ".join(foc)
    return (
        f"{role_title} — Data-driven problem-solver with a track record of making media/data operations measurably better. "
        f"Strengths: {foc_txt}. Uses the XYZ rule to surface impact and match hiring needs."
    )

def relevance_positioning_score(cv_text: str, jd_text: str):
    cv_tokens = set(clean_tokens(cv_text))
    jd_tokens = set(clean_tokens(jd_text))
    if not jd_tokens:
        return 0.0, {}
    overlap = len(cv_tokens & jd_tokens) / max(1, len(jd_tokens))
    seniority = 1.0 if (SENIORITY_TERMS & jd_tokens) else 0.7
    impact_bias = 0.9 if any(t in cv_tokens for t in IMPACT_TERMS) else 0.7
    score = (0.6*overlap + 0.25*seniority + 0.15*impact_bias) * 10
    # Explainers
    missing = sorted(list(jd_tokens - cv_tokens))[:15]
    present = sorted(list(cv_tokens & jd_tokens))[:15]
    details = {"present_keywords": present, "missing_keywords": missing}
    return round(min(10.0, score), 2), details

def save_version(title: str, tvps: str, sections: dict, jd_snippet: str):
    raw = write_docx(title, tvps, sections)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{re.sub(r'[^a-zA-Z0-9]+','-', title)[:60]}_{ts}.docx"
    path = SAVED_DIR / filename
    path.write_bytes(raw)
    meta = json.loads(META_FILE.read_text())
    meta.insert(0, {
        "title": title,
        "file": filename,
        "saved_at": ts,
        "jd_excerpt": jd_snippet[:280]
    })
    META_FILE.write_text(json.dumps(meta, indent=2))
    return path

def load_history():
    if META_FILE.exists():
        return json.loads(META_FILE.read_text())
    return []

def parse_bullets_from_text(text: str):
    lines = [l.strip("•- \t") for l in text.splitlines() if l.strip()]
    # Retain only bullet-like lines (heuristic)
    bullets = []
    for l in lines:
        if len(l.split()) >= 3:
            bullets.append(l.rstrip("."))
    return bullets

# ----------------------------- #
# UI                            #
# ----------------------------- #

st.set_page_config(page_title="CV Tailor • XYZ, Quant, Save", page_icon="🧰", layout="wide")
st.title("🧰 CV Tailor — JD‑aware, XYZ rewrites, and reusable bases")

with st.sidebar:
    st.header("⚙️ Settings")
    default_role = st.text_input("Target Role Title", value="Strategy & Transformation / Data Governance Manager")
    section_name = st.text_input("Main Experience Section Name", value="Experience")
    st.markdown("---")
    st.caption("Tip: Keep a clean base CV in .docx or paste bullets as text.")

col1, col2 = st.columns(2, gap="large")

with col1:
    st.subheader("1) Paste Job Description")
    jd_text = st.text_area("Job Description", height=280, placeholder="Paste the full JD here...")
    st.subheader("2) Provide Your Current CV")
    uploaded = st.file_uploader("Upload base CV (.docx) or paste bullets below", type=["docx"])
    cv_text = ""
    base_bullets = []
    if uploaded is not None:
        try:
            cv_text = read_docx(uploaded.read())
            st.success("Base CV loaded from .docx")
        except Exception as e:
            st.error(f"Couldn't read .docx: {e}")
    manual = st.text_area("Or paste bullet points", height=220, placeholder="• Drove X in Y\n• Improved Z by 12%\n...")
    if manual.strip():
        cv_text += "\n" + manual
    if cv_text.strip():
        base_bullets = parse_bullets_from_text(cv_text)

with col2:
    st.subheader("3) Hiring Manager POV")
    if jd_text.strip():
        hm = hiring_manager_notes(jd_text)
        st.markdown("**Must-haves:** " + ", ".join(hm["Must-haves"]) if hm["Must-haves"] else "_(none detected)_")
        st.markdown("**Great-to-haves:** " + ", ".join(hm["Great-to-haves"]) if hm["Great-to-haves"] else "_(none detected)_")
        st.markdown("**Impact to Prove:** " + ", ".join(hm["Impact to Prove"]) if hm["Impact to Prove"] else "_(none detected)_")
    else:
        st.info("Paste the JD to generate manager POV guidance.")

    if cv_text.strip() and jd_text.strip():
        score, details = relevance_positioning_score(cv_text, jd_text)
        st.metric("Relevance & Positioning Score", f"{score}/10")
        with st.expander("Keyword overlap (present / missing)"):
            st.write("**Present:**", ", ".join(details["present_keywords"]) or "—")
            st.write("**Missing:**", ", ".join(details["missing_keywords"]) or "—")

st.markdown("---")
st.subheader("4) Auto‑rewrite bullets to the XYZ rule")

role_kws = top_keywords(jd_text, n=12) if jd_text.strip() else []
jd_outcomes = [k for k in role_kws if any(t in k for t in ["increase","reduce","improve","optimize","growth","scale"])]

if base_bullets:
    # Table of rewrites with quant prompts
    rewritten = []
    quant_inputs = []
    for i, b in enumerate(base_bullets, start=1):
        rw = xyz_rewrite(b, role_keywords=role_kws, jd_outcomes=jd_outcomes)
        with st.expander(f"Bullet {i}: {b[:80]}{'...' if len(b)>80 else ''}"):
            st.write("Suggested (XYZ):")
            new_text = st.text_area(f"Rewrite {i}", value=rw, key=f"rw_{i}")
            missing_metric = "⚠ add a measurable result" in new_text
            q = st.text_input(f"Optional: add a quant/impact for {i} (e.g., '+18%', '£250k', '2 days/wk saved')",
                              value="", key=f"q_{i}")
            if q and missing_metric:
                new_text = new_text.replace("— ⚠ add a measurable result (%, £, time saved, error rate, etc.)", f"by {q}")
            rewritten.append(new_text.strip())
            quant_inputs.append(q)
else:
    st.warning("Add your base CV (upload .docx or paste bullets) to generate rewrites.")
    rewritten = []

st.markdown("---")
st.subheader("5) Targeted Value Proposition Statement (top of CV)")
tvps = ""
if jd_text.strip():
    tvps = tvps_generator(default_role, jd_text)
tvps = st.text_area("TVPS", value=tvps, height=90)

st.markdown("---")
st.subheader("6) Package & Save")

final_sections = {section_name: rewritten}
cv_title = st.text_input("Output CV Title", value=f"{default_role} — Tailored CV")
if st.button("💾 Save tailored CV"):
    if not rewritten:
        st.error("No rewritten bullets yet.")
    else:
        path = save_version(cv_title, tvps, final_sections, jd_snippet=jd_text.strip())
        st.success(f"Saved: {path.name}")
        with open(path, "rb") as f:
            st.download_button("⬇️ Download .docx", data=f.read(), file_name=path.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("---")
st.subheader("📚 Saved Bases (reusable)")
history = load_history()
if history:
    df = pd.DataFrame(history)
    st.dataframe(df[["saved_at","title","file","jd_excerpt"]], use_container_width=True, hide_index=True)
    chosen = st.selectbox("Open a saved base", ["—"] + [h["file"] for h in history])
    if chosen != "—":
        file_path = SAVED_DIR / chosen
        st.info(f"Selected: {chosen}")
        with open(file_path, "rb") as f:
            st.download_button("⬇️ Download selected .docx", data=f.read(), file_name=chosen, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.caption("No saved CVs yet. When you save, they’ll appear here.")

st.markdown("---")
with st.expander("What this app optimises for"):
    st.markdown("""
- **Hiring manager POV:** surfaces must-haves, great-to-haves, and impact themes from the JD.
- **XYZ rule:** rewrites every bullet into *Accomplished X by doing Y, resulting in Z*; prompts you to add missing numbers.
- **Quantification-first:** flags bullets without measurable outcomes and lets you inject %/£/time metrics quickly.
- **Relevance & Positioning Score:** quick 0–10 score using keyword overlap + seniority/impact signals.
- **Reusable bases:** every approved version is saved as a .docx with metadata for future tailoring.
""")